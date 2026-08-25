import os
import sys
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

# Colors
C_PRIMARY = RGBColor(23, 35, 61)     # #17233D Navy
C_SECONDARY = RGBColor(8, 123, 113)  # #087B71 Teal
C_ACCENT = RGBColor(238, 158, 30)    # #EE9E1E Amber
C_BODY = RGBColor(45, 55, 72)        # #2D3748 Charcoal
C_MUTED = RGBColor(102, 112, 133)    # #667085 Slate
C_WHITE = RGBColor(255, 255, 255)

HEX_PRIMARY = "17233D"
HEX_SECONDARY = "087B71"
HEX_ACCENT = "EE9E1E"
HEX_LIGHT_BG = "F4F7F9"
HEX_CALLOUT_BG = "EDF7F5"
HEX_BORDER = "DFE5EA"
HEX_ALT_ROW = "F8FAFB"

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=140, right=140):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'''<w:tcMar {nsdecls("w")}>
        <w:top w:w="{top}" w:type="dxa"/>
        <w:bottom w:w="{bottom}" w:type="dxa"/>
        <w:left w:w="{left}" w:type="dxa"/>
        <w:right w:w="{right}" w:type="dxa"/>
    </w:tcMar>''')
    tcPr.append(tcMar)

def set_table_borders(table, color="D3D9DE", sz="4", val="single"):
    tblPr = table._tbl.tblPr
    tblBorders = parse_xml(f'''<w:tblBorders {nsdecls("w")}>
        <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
        <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
        <w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
        <w:insideV w:val="none"/>
        <w:left w:val="none"/>
        <w:right w:val="none"/>
    </w:tblBorders>''')
    tblPr.append(tblBorders)

def add_header_footer(doc):
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.different_first_page_header_footer = True
        
        # Header
        header = section.header
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hrun = hp.add_run("TRIJOTECH Website — Standard Operating Procedure (SOP)")
        hrun.font.name = "Poppins"
        hrun.font.size = Pt(8.5)
        hrun.font.color.rgb = C_MUTED
        
        # Footer
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        frun = fp.add_run("TRIJOTECH Confidential | Standard Operating Procedure & Technical Specification | Version 1.0")
        frun.font.name = "Poppins"
        frun.font.size = Pt(8)
        frun.font.color.rgb = C_MUTED

def style_heading(p, text, level=1):
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = "Poppins"
    run.bold = True
    if level == 1:
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)
        run.font.size = Pt(16)
        run.font.color.rgb = C_PRIMARY
    elif level == 2:
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(4)
        run.font.size = Pt(13)
        run.font.color.rgb = C_SECONDARY
    elif level == 3:
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(3)
        run.font.size = Pt(11)
        run.font.color.rgb = C_PRIMARY

def add_p(doc, text="", bold_prefix="", space_after=4, space_before=0, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.bold = True
        r_pre.font.name = "Poppins"
        r_pre.font.size = Pt(10)
        r_pre.font.color.rgb = C_PRIMARY
    if text:
        r = p.add_run(text)
        r.font.name = "Poppins"
        r.font.size = Pt(10)
        r.font.color.rgb = C_BODY
        r.italic = italic
    return p

def add_bullet(doc, text, bold_prefix="", level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.bold = True
        r_pre.font.name = "Poppins"
        r_pre.font.size = Pt(9.5)
        r_pre.font.color.rgb = C_PRIMARY
    if text:
        r = p.add_run(text)
        r.font.name = "Poppins"
        r.font.size = Pt(9.5)
        r.font.color.rgb = C_BODY
    return p

def add_callout(doc, text, title="IMPORTANT"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    cell = tbl.cell(0, 0)
    cell.width = Inches(6.5)
    
    border_color = HEX_SECONDARY if title in ["NOTE", "TIP", "BEST PRACTICE"] else HEX_ACCENT if title in ["IMPORTANT", "SECURITY"] else "C0392B"
    bg_color = HEX_CALLOUT_BG if title in ["NOTE", "TIP", "BEST PRACTICE"] else "FEF9E7" if title in ["IMPORTANT", "SECURITY"] else "FADBD8"
    
    set_cell_background(cell, bg_color)
    set_cell_margins(cell, top=120, bottom=120, left=160, right=160)
    
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'''<w:tcBorders {nsdecls("w")}>
        <w:top w:val="none"/>
        <w:left w:val="single" w:sz="24" w:space="0" w:color="{border_color}"/>
        <w:bottom w:val="none"/>
        <w:right w:val="none"/>
    </w:tcBorders>''')
    tcPr.append(tcBorders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    run_t = p.add_run(f"[{title}] ")
    run_t.bold = True
    run_t.font.name = "Poppins"
    run_t.font.size = Pt(9.5)
    run_t.font.color.rgb = C_SECONDARY if border_color == HEX_SECONDARY else C_ACCENT if border_color == HEX_ACCENT else RGBColor(192, 57, 43)
    
    run = p.add_run(text)
    run.font.name = "Poppins"
    run.font.size = Pt(9.5)
    run.font.color.rgb = C_BODY
    doc.add_paragraph().paragraph_format.space_after = Pt(3)

def add_code_block(doc, code_text):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    cell = tbl.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_background(cell, "1E293B")
    set_cell_margins(cell, top=100, bottom=100, left=140, right=140)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(code_text)
    run.font.name = "Consolas"
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(241, 245, 249)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)

def create_table(doc, headers, data, col_widths=None):
    tbl = doc.add_table(rows=len(data) + 1, cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    set_table_borders(tbl)
    
    # Header Row
    hdr_cells = tbl.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        set_cell_background(hdr_cells[i], HEX_PRIMARY)
        set_cell_margins(hdr_cells[i], top=120, bottom=120, left=120, right=120)
        p = hdr_cells[i].paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        for r in p.runs:
            r.font.name = "Poppins"
            r.font.size = Pt(9)
            r.bold = True
            r.font.color.rgb = C_WHITE
        if col_widths and i < len(col_widths):
            hdr_cells[i].width = Inches(col_widths[i])
            
    # Set header repeat across pages
    trPr = tbl.rows[0]._tr.get_or_add_trPr()
    trPr.append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))
    
    # Data Rows
    for r_idx, row_data in enumerate(data):
        row_cells = tbl.rows[r_idx + 1].cells
        bg = HEX_ALT_ROW if r_idx % 2 == 1 else "FFFFFF"
        for c_idx, val in enumerate(row_data):
            row_cells[c_idx].text = str(val)
            set_cell_background(row_cells[c_idx], bg)
            set_cell_margins(row_cells[c_idx], top=80, bottom=80, left=100, right=100)
            p = row_cells[c_idx].paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.1
            for r in p.runs:
                r.font.name = "Poppins"
                r.font.size = Pt(8.5)
                r.font.color.rgb = C_BODY
            if col_widths and c_idx < len(col_widths):
                row_cells[c_idx].width = Inches(col_widths[c_idx])
                
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return tbl

print("Base layout & styling initialized.")
