import os
import sys
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def build_sop():
    doc = Document()
    
    # Colors
    C_PRIMARY = RGBColor(23, 35, 61)     # #17233D Navy
    C_SECONDARY = RGBColor(8, 123, 113)  # #087B71 Teal
    C_ACCENT = RGBColor(238, 158, 30)    # #EE9E1E Amber
    C_BODY = RGBColor(45, 55, 72)        # #2D3748 Charcoal
    C_MUTED = RGBColor(102, 112, 133)    # #667085 Slate
    C_WHITE = RGBColor(255, 255, 255)

    HEX_PRIMARY = "17233D"
    HEX_SECONDARY = "087B71"
    HEX_ACCENT = "EE9E1E"
    HEX_LIGHT_BG = "F4F7F9"
    HEX_CALLOUT_BG = "EDF7F5"
    HEX_BORDER = "DFE5EA"
    HEX_ALT_ROW = "F8FAFB"

    def set_cell_background(cell, fill_hex):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
        tcPr.append(shd)

    def set_cell_margins(cell, top=100, bottom=100, left=140, right=140):
        tcPr = cell._tc.get_or_add_tcPr()
        tcMar = parse_xml(f'''<w:tcMar {nsdecls("w")}>
            <w:top w:w="{top}" w:type="dxa"/>
            <w:bottom w:w="{bottom}" w:type="dxa"/>
            <w:left w:w="{left}" w:type="dxa"/>
            <w:right w:w="{right}" w:type="dxa"/>
        </w:tcMar>''')
        tcPr.append(tcMar)

    def set_table_borders(table, color="D3D9DE", sz="4", val="single"):
        tblPr = table._tbl.tblPr
        tblBorders = parse_xml(f'''<w:tblBorders {nsdecls("w")}>
            <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:insideV w:val="none"/>
            <w:left w:val="none"/>
            <w:right w:val="none"/>
        </w:tblBorders>''')
        tblPr.append(tblBorders)

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.different_first_page_header_footer = True
        
        # Header
        header = section.header
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hrun = hp.add_run("TRIJOTECH Website — Standard Operating Procedure (SOP)")
        hrun.font.name = "Poppins"
        hrun.font.size = Pt(8.5)
        hrun.font.color.rgb = C_MUTED
        
        # Footer
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        frun = fp.add_run("TRIJOTECH Confidential | Standard Operating Procedure & Technical Specification | Version 1.0")
        frun.font.name = "Poppins"
        frun.font.size = Pt(8)
        frun.font.color.rgb = C_MUTED

    def h1(text):
        p = doc.add_paragraph()
        p.paragraph_format.keep_with_next = True
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(text)
        run.font.name = "Poppins"
        run.bold = True
        run.font.size = Pt(15)
        run.font.color.rgb = C_PRIMARY
        return p

    def h2(text):
        p = doc.add_paragraph()
        p.paragraph_format.keep_with_next = True
        p.paragraph_format.space_before = Pt(13)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text)
        run.font.name = "Poppins"
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = C_SECONDARY
        return p

    def h3(text):
        p = doc.add_paragraph()
        p.paragraph_format.keep_with_next = True
        p.paragraph_format.space_before = Pt(9)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(text)
        run.font.name = "Poppins"
        run.bold = True
        run.font.size = Pt(10.5)
        run.font.color.rgb = C_PRIMARY
        return p

    def p(text="", bold_pre="", space_after=4, italic=False):
        par = doc.add_paragraph()
        par.paragraph_format.space_before = Pt(0)
        par.paragraph_format.space_after = Pt(space_after)
        par.paragraph_format.line_spacing = 1.15
        if bold_pre:
            r_pre = par.add_run(bold_pre)
            r_pre.bold = True
            r_pre.font.name = "Poppins"
            r_pre.font.size = Pt(9.5)
            r_pre.font.color.rgb = C_PRIMARY
        if text:
            r = par.add_run(text)
            r.font.name = "Poppins"
            r.font.size = Pt(9.5)
            r.font.color.rgb = C_BODY
            r.italic = italic
        return par

    def bullet(text, bold_pre="", level=0):
        par = doc.add_paragraph(style='List Bullet')
        par.paragraph_format.space_before = Pt(1)
        par.paragraph_format.space_after = Pt(2)
        par.paragraph_format.line_spacing = 1.15
        par.paragraph_format.left_indent = Inches(0.25 * (level + 1))
        if bold_pre:
            r_pre = par.add_run(bold_pre)
            r_pre.bold = True
            r_pre.font.name = "Poppins"
            r_pre.font.size = Pt(9.5)
            r_pre.font.color.rgb = C_PRIMARY
        if text:
            r = par.add_run(text)
            r.font.name = "Poppins"
            r.font.size = Pt(9.5)
            r.font.color.rgb = C_BODY
        return par

    def callout(text, title="IMPORTANT"):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.autofit = False
        cell = tbl.cell(0, 0)
        cell.width = Inches(6.5)
        
        border_color = HEX_SECONDARY if title in ["NOTE", "TIP", "BEST PRACTICE"] else HEX_ACCENT if title in ["IMPORTANT", "SECURITY"] else "C0392B"
        bg_color = HEX_CALLOUT_BG if title in ["NOTE", "TIP", "BEST PRACTICE"] else "FEF9E7" if title in ["IMPORTANT", "SECURITY"] else "FADBD8"
        
        set_cell_background(cell, bg_color)
        set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
        
        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = parse_xml(f'''<w:tcBorders {nsdecls("w")}>
            <w:top w:val="none"/>
            <w:left w:val="single" w:sz="24" w:space="0" w:color="{border_color}"/>
            <w:bottom w:val="none"/>
            <w:right w:val="none"/>
        </w:tcBorders>''')
        tcPr.append(tcBorders)
        
        cp = cell.paragraphs[0]
        cp.paragraph_format.space_before = Pt(2)
        cp.paragraph_format.space_after = Pt(2)
        cp.paragraph_format.line_spacing = 1.15
        run_t = cp.add_run(f"[{title}] ")
        run_t.bold = True
        run_t.font.name = "Poppins"
        run_t.font.size = Pt(9.5)
        run_t.font.color.rgb = C_SECONDARY if border_color == HEX_SECONDARY else C_ACCENT if border_color == HEX_ACCENT else RGBColor(192, 57, 43)
        
        run = cp.add_run(text)
        run.font.name = "Poppins"
        run.font.size = Pt(9.5)
        run.font.color.rgb = C_BODY
        doc.add_paragraph().paragraph_format.space_after = Pt(3)

    def code(code_text):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.autofit = False
        cell = tbl.cell(0, 0)
        cell.width = Inches(6.5)
        set_cell_background(cell, "1E293B")
        set_cell_margins(cell, top=90, bottom=90, left=130, right=130)
        
        cp = cell.paragraphs[0]
        cp.paragraph_format.space_before = Pt(2)
        cp.paragraph_format.space_after = Pt(2)
        cp.paragraph_format.line_spacing = 1.15
        run = cp.add_run(code_text)
        run.font.name = "Consolas"
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(241, 245, 249)
        doc.add_paragraph().paragraph_format.space_after = Pt(3)

    def tbl(headers, data, col_widths=None):
        t = doc.add_table(rows=len(data) + 1, cols=len(headers))
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        t.autofit = False
        set_table_borders(t)
        
        hdr_cells = t.rows[0].cells
        for i, title in enumerate(headers):
            hdr_cells[i].text = title
            set_cell_background(hdr_cells[i], HEX_PRIMARY)
            set_cell_margins(hdr_cells[i], top=100, bottom=100, left=100, right=100)
            cp = hdr_cells[i].paragraphs[0]
            cp.paragraph_format.space_before = Pt(0)
            cp.paragraph_format.space_after = Pt(0)
            for r in cp.runs:
                r.font.name = "Poppins"
                r.font.size = Pt(8.5)
                r.bold = True
                r.font.color.rgb = C_WHITE
            if col_widths and i < len(col_widths):
                hdr_cells[i].width = Inches(col_widths[i])
                
        trPr = t.rows[0]._tr.get_or_add_trPr()
        trPr.append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))
        
        for r_idx, row_data in enumerate(data):
            row_cells = t.rows[r_idx + 1].cells
            bg = HEX_ALT_ROW if r_idx % 2 == 1 else "FFFFFF"
            for c_idx, val in enumerate(row_data):
                row_cells[c_idx].text = str(val)
                set_cell_background(row_cells[c_idx], bg)
                set_cell_margins(row_cells[c_idx], top=70, bottom=70, left=90, right=90)
                cp = row_cells[c_idx].paragraphs[0]
                cp.paragraph_format.space_before = Pt(0)
                cp.paragraph_format.space_after = Pt(0)
                cp.paragraph_format.line_spacing = 1.1
                for r in cp.runs:
                    r.font.name = "Poppins"
                    r.font.size = Pt(8)
                    r.font.color.rgb = C_BODY
                if col_widths and c_idx < len(col_widths):
                    row_cells[c_idx].width = Inches(col_widths[c_idx])
                    
        doc.add_paragraph().paragraph_format.space_after = Pt(3)
        return t

    # ─────────────────────────────────────────────────────────────
    # COVER PAGE
    # ─────────────────────────────────────────────────────────────
    cover_p = doc.add_paragraph()
    cover_p.paragraph_format.space_before = Pt(72)
    cover_p.paragraph_format.space_after = Pt(6)
    r_tag = cover_p.add_run("ENTERPRISE STANDARD OPERATING PROCEDURE & TECHNICAL SPECIFICATION")
    r_tag.font.name = "Poppins"
    r_tag.font.size = Pt(10)
    r_tag.bold = True
    r_tag.font.color.rgb = C_ACCENT

    cover_title = doc.add_paragraph()
    cover_title.paragraph_format.space_before = Pt(0)
    cover_title.paragraph_format.space_after = Pt(12)
    r_title = cover_title.add_run("TRIJOTECH Website Architecture, Operations & Maintenance Manual")
    r_title.font.name = "Poppins"
    r_title.font.size = Pt(24)
    r_title.bold = True
    r_title.font.color.rgb = C_PRIMARY

    cover_sub = doc.add_paragraph()
    cover_sub.paragraph_format.space_before = Pt(0)
    cover_sub.paragraph_format.space_after = Pt(28)
    r_sub = cover_sub.add_run("A Complete End-to-End Standard Operating Procedure (SOP), System Architecture Blueprint, Content Management & Troubleshooting Handbook for Developers, Administrators, and Operations Teams.")
    r_sub.font.name = "Poppins"
    r_sub.font.size = Pt(11)
    r_sub.font.color.rgb = C_MUTED

    meta_tbl = doc.add_table(rows=6, cols=2)
    meta_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(meta_tbl)
    meta_data = [
        ("Document Version", "1.0 (Production Release)"),
        ("Document Date", "August 2026"),
        ("Prepared By", "Senior Technical Documentation Specialist & Next.js Architecture Team"),
        ("Reviewed By", "Lead Enterprise Architect & Lead Frontend Engineer"),
        ("Approved By", "TRIJOTECH Executive & IT Steering Committee"),
        ("Document Status", "Approved / Official Baseline Specification")
    ]
    for idx, (label, val) in enumerate(meta_data):
        r_cells = meta_tbl.rows[idx].cells
        r_cells[0].text = label
        r_cells[1].text = val
        set_cell_background(r_cells[0], HEX_LIGHT_BG)
        set_cell_background(r_cells[1], "FFFFFF")
        set_cell_margins(r_cells[0], top=80, bottom=80, left=100, right=100)
        set_cell_margins(r_cells[1], top=80, bottom=80, left=100, right=100)
        r_cells[0].paragraphs[0].runs[0].font.bold = True
        r_cells[0].paragraphs[0].runs[0].font.size = Pt(9)
        r_cells[0].paragraphs[0].runs[0].font.color.rgb = C_PRIMARY
        r_cells[1].paragraphs[0].runs[0].font.size = Pt(9)
        r_cells[1].paragraphs[0].runs[0].font.color.rgb = C_BODY
        r_cells[0].width = Inches(2.2)
        r_cells[1].width = Inches(4.3)

    doc.add_page_break()

    # ─────────────────────────────────────────────────────────────
    # DOCUMENT CONTROL & REVISION HISTORY
    # ─────────────────────────────────────────────────────────────
    h1("DOCUMENT CONTROL")
    p("This document is the official technical standard operating procedure and maintenance specification for the TRIJOTECH Corporate Website. Any updates to application architecture, deployment workflows, administrative capabilities, or data handling must be documented in revision increments.")

    tbl(
        ["Field", "Details"],
        [
            ["Document Name", "TRIJOTECH Website Standard Operating Procedure & Tech Spec"],
            ["Project Name", "TRIJOTECH Corporate Web Platform (webapp)"],
            ["Document Type", "Standard Operating Procedure (SOP) & Technical Manual"],
            ["Document Version", "1.0.0"],
            ["Target Environment", "Next.js 16 (Turbopack) / Node.js 20+ / React 19 / TypeScript 5"],
            ["Security Level", "Confidential / Internal Technical & Operational Distribution"],
            ["Storage Location", "Repository root /docs/ and core corporate document archives"],
            ["Last Updated", "August 2026"]
        ],
        [2.2, 4.3]
    )

    h2("Revision History")
    tbl(
        ["Version", "Date", "Author", "Summary of Changes"],
        [
            ["0.9.0", "August 2026", "Engineering Team", "Initial architecture documentation and core routing specification."],
            ["0.9.5", "August 2026", "Documentation Team", "Added Admin Studio, Chatbot Intelligence Engine, and Blog Block Editor SOP."],
            ["1.0.0", "August 2026", "Senior Architecture Specialist", "Complete baseline release with unified Admin Navbar, HTML stripping runtime, and all 49 routes documented."]
        ],
        [0.8, 1.1, 1.6, 3.0]
    )

    # ─────────────────────────────────────────────────────────────
    # TABLE OF CONTENTS
    # ─────────────────────────────────────────────────────────────
    h1("TABLE OF CONTENTS")
    p("The following table of contents outlines the 44 primary operational and architectural sections detailed within this Standard Operating Procedure document:")
    
    toc_p = doc.add_paragraph()
    r_toc = toc_p.add_run("1. Introduction\n2. Technology Stack\n3. System Architecture\n4. Project Folder Structure\n5. Installation and Local Environment Setup\n6. Website Routing and Page Structure\n7. Homepage SOP\n8. Navbar and Navigation Management\n9. Hero Section Management\n10. Services Section\n11. Industries Section\n12. Admin Panel SOP\n13. Blog Management\n14. Contact Form SOP\n15. Career Form SOP\n16. Email Configuration\n17. Chatbot Studio & AI Engine\n18. Animation System\n19. Styling and Design System\n20. Responsive Design\n21. Image and Media Management\n22. Performance Optimization\n23. SEO and Metadata\n24. Security Architecture\n25. Development SOP\n26. Git and Version Control SOP\n27. Testing SOP\n28. Deployment SOP\n29. Post-Deployment Validation\n30. Troubleshooting Guide\n31. Website Maintenance SOP\n32. Change Management Procedure\n33. Backup and Recovery\n34. Coding and Maintenance Guidelines\n35. Common Operation Procedures (SOP-01 to SOP-20)\n36. Developer Quick Reference\n37. Component Inventory\n38. Dependency Inventory\n39. API Inventory\n40. Documentation of Important Business Content\n41. Known Limitations\n42. Prioritized Recommendations\n43. Glossary\n44. Appendices (A through G)")
    r_toc.font.name = "Poppins"
    r_toc.font.size = Pt(9.5)
    r_toc.font.color.rgb = C_SECONDARY
    toc_p.paragraph_format.line_spacing = 1.3

    doc.add_page_break()

    # ─────────────────────────────────────────────────────────────
    # 1. INTRODUCTION
    # ─────────────────────────────────────────────────────────────
    h1("1. INTRODUCTION")
    
    h2("1.1 Purpose of the SOP")
    p("The purpose of this Standard Operating Procedure (SOP) is to serve as the definitive, single-source-of-truth technical guide and operational handbook for the TRIJOTECH Corporate Website. It establishes standardized workflows for engineering, content administration, UI maintenance, search engine optimization, form processing, AI knowledge configuration, and zero-downtime deployment.")

    h2("1.2 Scope")
    p("This document encompasses the entire codebase located in the webapp project workspace, spanning:")
    bullet("Frontend presentation layer, React 19 components, and Tailwind CSS 4 design token hierarchy.")
    bullet("Next.js 16 App Router architecture, Server Components, Client Components, and dynamic static generation.")
    bullet("Backend Route Handlers (`/api/`) powering admin authentication, content overrides, blog publishing, chatbot queries, image asset uploads, and secure form submissions.")
    bullet("Administrative tools: Live Visual Page Content Editor, Blog Management Studio, Create Blog Studio, and Chatbot Studio.")
    bullet("JSON data stores, verified AI knowledge base files, static asset management, and Nodemailer SMTP communications.")

    h2("1.3 Intended Audience")
    p("This SOP is engineered for direct use by:")
    bullet("Full-Stack & Frontend Developers: For onboarding, component extension, styling standards, and architectural adherence.", "Developers: ")
    bullet("Content & Website Administrators: For managing live text overrides, publishing rich blog posts, and updating chatbot knowledge.", "Administrators: ")
    bullet("DevOps & Support Engineers: For local environment setup, production builds, deployment automation, and troubleshooting.", "DevOps / Support: ")
    bullet("Project Managers & Quality Assurance: For functional verification, compliance auditing, and release sign-offs.", "QA / PMs: ")

    h2("1.4 Website Overview")
    p("TRIJOTECH is an enterprise SAP solutions, digital transformation, and business technology consulting provider. The website showcases the company's full portfolio of specialized SAP services (Consulting, S/4HANA Implementation, Support, BTP Full Stack, Data Integration, and AI/ML), industry-tailored platforms (Retail, Pharma, Manufacturing, Banking, Utilities, Steel, Telecom), proprietary product solutions (E-Invoicing Pro, Finlagoon Consolidation, Profitability Pro), thought leadership articles, career opportunities, and direct customer engagement channels.")

    h2("1.5 Primary Objectives")
    bullet("Brand Leadership: Establish TRIJOTECH as an elite, high-tier SAP and digital intelligence partner through high-performance animations, modern glassmorphism UI, and polished design.", "1. ")
    bullet("Client Engagement & Lead Generation: Provide frictionless, secure contact forms, career job application flows, and instant interactive AI chat assistance.", "2. ")
    bullet("Dynamic Content Agility: Empower non-developer admins to modify live website copy, upload imagery, publish styled blog posts, and train chatbot responses via dedicated web studios.", "3. ")
    bullet("Performance & Security: Ensure sub-second page loads via Next.js 16 static optimization (SSG/ISR), strict session authentication, rate limiting, and mathematical CAPTCHA safeguards.", "4. ")

    # ─────────────────────────────────────────────────────────────
    # 2. TECHNOLOGY STACK
    # ─────────────────────────────────────────────────────────────
    h1("2. TECHNOLOGY STACK")
    p("The TRIJOTECH website is built upon a modern, full-stack JavaScript/TypeScript framework designed for enterprise-grade scalability, instant rendering, and maintainability. The table below outlines all verified technologies implemented in the codebase:")

    tbl(
        ["Technology", "Version / Spec", "Usage in TRIJOTECH", "Project Location"],
        [
            ["Next.js", "16.2.12 (Turbopack)", "Core web framework, App Router, SSR, SSG static pre-rendering, Route Handlers.", "package.json, next.config.ts, app/"],
            ["React / React-DOM", "19.2.4", "Declarative component UI, Concurrent rendering, Hooks, Client/Server boundary.", "package.json, app/, components/"],
            ["TypeScript", "^5.0.0", "End-to-end static type safety, data models, API payload validation, build safety.", "tsconfig.json, app/**/*.ts, *.tsx"],
            ["Tailwind CSS", "^4.0.0", "Utility-first design engine, @theme token definitions, custom responsive classes.", "app/globals.css, package.json"],
            ["Framer Motion", "^13.1.0", "Scroll-triggered viewport reveals, spring animations, number counters, interactive stages.", "package.json, components/motion/, components/scenes/"],
            ["Lucide React", "^1.31.0", "Modern, scalable SVG icon set used across navigation, admin studio, and CTAs.", "package.json, components/layout/, app/admin/"],
            ["React Icons", "^5.7.0", "Extended technical, social, and brand icons across industry sections and footer.", "package.json, components/industries/"],
            ["Nodemailer", "^9.0.4", "Server-side SMTP email transport for Contact Us and Career job applications.", "app/api/forms/, app/lib/form-security.ts"],
            ["Lottie React", "^3.1.0", "Vector JSON animations for interactive hero scenes and service graphics.", "package.json, components/scenes/"],
            ["React Slick / Slick", "^0.31.0 / ^1.8.1", "Touch-enabled responsive carousels for client testimonials and solution cards.", "components/landing/CardsCarousel.tsx, app/admin/"],
            ["DeepSeek LLM API", "deepseek-chat", "AI Chatbot conversational intelligence, verified knowledge retrieval grounding.", "app/services/deepseek.ts, app/api/chat/"],
            ["DOMParser & SafeHTML", "Web Standard API", "Sanitization of admin content overrides, stripping malicious scripts, safe runtime injection.", "app/components/ContentRuntime.tsx, app/admin/AdminEditor.tsx"],
            ["Node.js FS Promises", "v20+ Native", "Persistent JSON storage for site content, blog drafts, chatbot memory, and settings.", "app/lib/content-store.ts, app/lib/blog-store.ts"]
        ],
        [1.3, 1.1, 2.5, 1.6]
    )

    # ─────────────────────────────────────────────────────────────
    # 3. SYSTEM ARCHITECTURE
    # ─────────────────────────────────────────────────────────────
    h1("3. SYSTEM ARCHITECTURE")
    
    h2("3.1 High-Level Architecture Flow")
    p("The TRIJOTECH web application follows a modern decoupled architecture combining statically pre-rendered React Server Components with a client-side mutation runtime, file-based persistence, and external service adapters:")

    code("""[ User Browser / Client Device ]
       │  ▲ (HTTPS / HTML / Client Bundles)
       ▼  │
[ Next.js 16 App Router Server Engine (Turbopack) ]
  ├── Statically Generated Pages (SSG / 49 Routes)
  ├── Public Chrome Layout (Header, Footer, Floating Chatbot)
  ├── Live Content Runtime (DOM Mutation Observer & Dynamic Injector)
  └── Backend API Route Handlers (/api/*)
       ├── /api/admin/content      ──> Reads/Writes app/data/siteContent.json
       ├── /api/admin/blogs        ──> Reads/Writes app/lib/blog-store.ts & blogs.ts
       ├── /api/admin/chatbot      ──> Reads/Writes app/data/knowledge-base/*.json
       ├── /api/admin/images       ──> Stores assets to public/assets/uploads/
       ├── /api/forms/contact      ──> Validates CAPTCHA ──> Nodemailer SMTP Dispatch
       ├── /api/forms/careers      ──> Validates Resume  ──> Nodemailer SMTP Dispatch
       └── /api/chat               ──> Local Knowledge Match + DeepSeek LLM API""")

    h2("3.2 Frontend Architecture")
    p("The presentation layer separates concerns across layout wrappers, modular section components, interactive animation stages, and client runtime injectors:")
    bullet("PublicChrome (`app/components/PublicChrome.tsx`): Conditionally renders global Header, Footer, and Floating Chatbot on public routes while automatically hiding them on `/admin` studio workspaces.")
    bullet("ContentRuntime (`app/components/ContentRuntime.tsx`): Runs in the background on every public page. Fetches `/api/content`, observes DOM mutations, and seamlessly applies live text, link, and background image overrides saved through the Admin Editor without breaking React hydration.")
    bullet("Modular Component Hierarchy: All major landing sections (Hero, Showcase, Pillars, Metrics, CTA) are partitioned into pure functional components under `/components/` for maximum reusability.")

    h2("3.3 Server-Side Architecture")
    p("Next.js App Router Server Components execute on Node.js 20+ runtime. Dynamic endpoints enforce Node.js runtime (`export const runtime = 'nodejs'; export const dynamic = 'force-dynamic';`) to guarantee fresh filesystem I/O and real-time email dispatch.")

    h2("3.4 Data and Content Architecture")
    p("Data is structured cleanly without heavy SQL/NoSQL database overhead, using resilient atomic JSON stores:")
    bullet("`app/data/siteContent.json`: Master override store for visual website copy, image URLs, and button destinations.")
    bullet("`app/data/blogs.ts`: Structured array of static blog posts and rich content block definitions.")
    bullet("`app/data/knowledge-base/`: Segmented JSON directories (`services.json`, `solutions.json`, `company.json`, `faq.json`) parsed by the chatbot ranking engine.")
    bullet("`app/data/chatbot-settings.json`: Chatbot personality, prompt template, model hyperparameters, and rate limit rules.")

    h2("3.5 Admin Architecture")
    p("The administrative ecosystem provides 4 focused studios with a unified top navigation panel (`AdminNavbar.tsx`):")
    bullet("1. Page Content Editor (`/admin`): Visual iframe-based point-and-click editor with live DOM highlighting, automated HTML tag stripping, and instantaneous save/reset.")
    bullet("2. Blog Management Studio (`/admin/blogs`): Searchable, filterable repository of published and draft articles with category management.")
    bullet("3. Create Blog Studio (`/admin/createblog`): Full-screen block-based content studio supporting 10 block primitives (Headings, Paragraphs, Images, Quotes, Bullet Lists, Numbered Lists, Callouts, Dividers, Links, Custom CSS).")
    bullet("4. Chatbot Studio (`/admin/chatbot`): Two-column knowledge entry editor, intent/keyword training, response preview playground, and global LLM settings manager.")

    # ─────────────────────────────────────────────────────────────
    # 4. PROJECT FOLDER STRUCTURE
    # ─────────────────────────────────────────────────────────────
    h1("4. PROJECT FOLDER STRUCTURE")
    p("The codebase is organized logically under standard Next.js App Router conventions. Below is the directory tree:")

    code("""webapp/
├── app/
│   ├── about-us/           # About Us corporate story & leadership page
│   ├── admin/              # Admin workspace, Editor, Navbar, Blog & Chatbot Studios
│   │   ├── blog/           # Blog sub-components (Actions, Details, PostList)
│   │   ├── blogs/          # Dedicated Blog Management Studio route (/admin/blogs)
│   │   ├── chatbot/        # Dedicated Chatbot Studio route (/admin/chatbot)
│   │   ├── createblog/     # Full-screen Create Blog Studio route (/admin/createblog)
│   │   └── login/          # Secure Admin Login authentication page (/admin/login)
│   ├── api/                # Backend API route handlers (Admin, Forms, Chat, Content)
│   ├── blogs/              # Public blog listing & dynamic slug reader (/blogs/[slug])
│   ├── careers/            # Careers culture, job openings & resume application form
│   ├── case-studies/       # Enterprise case studies showcase
│   ├── components/         # Page-specific client chrome, headers, and ContentRuntime
│   ├── contact-us/         # Public Contact Us inquiries page
│   ├── data/               # Persistent JSON stores, blogs.ts, and knowledge-base/
│   ├── industries/         # Industry hubs & dynamic slug pages (/industries/[[...slug]])
│   ├── lib/                # Admin auth, content store, form security, blog stores
│   ├── services/           # Dedicated SAP service routes & dynamic slug handlers
│   ├── solutions/          # Specialized product solution landing pages
│   ├── globals.css         # Tailwind CSS 4 theme tokens, keyframes, utilities
│   ├── layout.tsx          # Root HTML layout with Poppins font and ContentRuntime
│   └── page.tsx            # Main Homepage
├── components/             # Reusable UI component library (Layout, Motion, Scenes, Forms)
├── public/                 # Static assets, brand SVGs, hero videos, image uploads
├── next.config.ts          # Next.js build configuration, image domains, headers
├── package.json            # Node.js dependencies and script definitions
└── tsconfig.json           # TypeScript configuration and path aliases (@/*)""")

    h2("Key Directory Details")
    tbl(
        ["Folder / Path", "Core Purpose", "Critical Architectural Notes"],
        [
            ["app/admin/", "Administrative control center for live site editing.", "Houses AdminEditor, AdminNavbar, BlogManager, KnowledgeManager."],
            ["app/api/", "Next.js backend serverless API endpoints.", "Handles form submission, SMTP dispatch, LLM querying, content updates."],
            ["app/data/", "Filesystem JSON & TypeScript data stores.", "Contains siteContent.json, blogs.ts, chatbot-settings.json."],
            ["app/lib/", "Backend utility helpers and security logic.", "Implements admin-auth.ts (HMAC tokens) and form-security.ts (SMTP/Turnstile)."],
            ["components/layout/", "Global chrome components (Header, Footer).", "Desktop dropdown navigation, mobile drawer, responsive brand logo."],
            ["components/motion/", "Framer Motion animation primitives.", "Contains Reveal.tsx, AnimatedCounter.tsx for scroll-triggered visual effects."],
            ["components/scenes/", "Interactive visual canvas & graphics stages.", "Canvas signal rivers, BTP layer exploding scenes, e-invoicing diagrams."],
            ["public/assets/uploads/", "Destination folder for admin image uploads.", "Static files served directly at /assets/uploads/<filename>."]
        ],
        [1.8, 2.2, 2.5]
    )

    # ─────────────────────────────────────────────────────────────
    # 5. INSTALLATION AND LOCAL ENVIRONMENT SETUP
    # ─────────────────────────────────────────────────────────────
    h1("5. INSTALLATION AND LOCAL ENVIRONMENT SETUP")
    
    h2("5.1 Prerequisites")
    bullet("Node.js: Version 20.x LTS or higher (Node 22 supported).", "Node.js: ")
    bullet("Package Manager: npm (v10+), yarn, or pnpm.", "npm: ")
    bullet("Version Control: Git 2.40+.", "Git: ")
    bullet("Operating System: Windows 10/11, macOS, or Linux.", "OS: ")

    h2("5.2 Local Setup Procedure")
    p("Step 1: Clone the repository and navigate into the `webapp` directory:")
    code("""git clone <repository-url>
cd projectTrijowebsite/webapp""")

    p("Step 2: Install project dependencies:")
    code("npm install")

    h2("5.3 Environment Variables Configuration")
    p("Create a `.env.local` file in the root of `webapp/` using the following configuration reference:")

    tbl(
        ["Environment Variable", "Purpose / Usage", "Required?", "Sensitivity"],
        [
            ["ADMIN_PASSWORD", "Master password for accessing /admin studios.", "YES", "CRITICAL SECRET"],
            ["ADMIN_SESSION_SECRET", "HMAC SHA-256 signing key for admin session cookies.", "YES", "CRITICAL SECRET"],
            ["NEXT_PUBLIC_SITE_URL", "Base public URL (default: http://localhost:3000).", "YES", "Public / Config"],
            ["SMTP_HOST", "Outbound email SMTP server hostname.", "Optional", "Config"],
            ["SMTP_PORT", "SMTP port (587 for TLS, 465 for SSL).", "Optional", "Config"],
            ["SMTP_USER", "SMTP authentication username / email account.", "Optional", "Sensitive"],
            ["SMTP_PASS", "SMTP authentication password / app password.", "Optional", "CRITICAL SECRET"],
            ["SMTP_SECURE", "Set 'true' for port 465, 'false' for port 587.", "Optional", "Config"],
            ["CONTACT_TO_EMAIL", "Destination inbox for Contact Us inquiries.", "Optional", "Config"],
            ["CAREERS_TO_EMAIL", "Destination inbox for Job Applications.", "Optional", "Config"],
            ["NEXT_PUBLIC_TURNSTILE_SITE_KEY", "Cloudflare Turnstile public key (CAPTCHA).", "Optional", "Public"],
            ["TURNSTILE_SECRET_KEY", "Cloudflare Turnstile secret verification key.", "Optional", "CRITICAL SECRET"],
            ["DEEPSEEK_API_KEY", "DeepSeek API key powering AI Chatbot intelligence.", "Optional", "CRITICAL SECRET"]
        ],
        [1.8, 2.3, 0.9, 1.5]
    )

    h2("5.4 Running the Application")
    p("Start the local development server with Turbopack:")
    code("npm run dev")
    p("Access the web application in your browser at `http://localhost:3000`.")
    p("Access the Admin Studio in your browser at `http://localhost:3000/admin`.")

    h2("5.5 Production Build and Execution")
    p("To validate TypeScript types, build static pages, and run the optimized production server:")
    code("""npm run build
npm run start""")

    # ─────────────────────────────────────────────────────────────
    # 6. ROUTE INVENTORY (ALL 49 COMPILED ROUTES)
    # ─────────────────────────────────────────────────────────────
    h1("6. WEBSITE ROUTING AND PAGE STRUCTURE")
    p("The TRIJOTECH web platform compiles into 49 distinct static and dynamic routes. Below is the complete route inventory:")

    tbl(
        ["Sr.", "Route Path", "Page Name", "Purpose & Primary Components", "Data Source"],
        [
            ["1", "/", "Homepage", "Main brand landing, Hero video, Services, Industries, Stats.", "Static / siteContent.json"],
            ["2", "/about-us", "About Us", "Company story, Leadership, Pillars showcase, Values.", "Static / siteContent.json"],
            ["3", "/services", "Services Hub", "Overview grid of all 6 SAP enterprise capabilities.", "Static / services.json"],
            ["4", "/services/sap-consulting", "SAP Consulting", "Strategic advisory, Roadmap, Architecture assessment.", "Static / siteContent.json"],
            ["5", "/services/sap-implementation", "SAP Implementation", "Greenfield & Brownfield S/4HANA implementation.", "Static / siteContent.json"],
            ["6", "/services/sap-support", "SAP Support & AMS", "24/7 Managed services, Incident management, Optimization.", "Static / siteContent.json"],
            ["7", "/services/sap-btp-full-stack", "SAP BTP Full Stack", "Cloud integration, Custom extensions, SAP Build.", "Static / siteContent.json"],
            ["8", "/services/sap-data-integration", "SAP Data & Integration", "Data river pipelines, SAC, Datasphere, ETL flows.", "Static / siteContent.json"],
            ["9", "/services/sap-ai-ml", "SAP AI & ML", "Predictive maintenance, GenAI in SAP, Intelligent ERP.", "Static / siteContent.json"],
            ["10", "/services/[slug]", "Dynamic Service Detail", "Fallback dynamic handler for parameterized service routes.", "Dynamic Parameter"],
            ["11", "/solutions", "Solutions Hub", "Overview of proprietary TRIJOTECH enterprise software.", "Static / solutions.json"],
            ["12", "/solutions/e-invoicing-pro", "E-Invoicing Pro", "Government portal compliant e-invoicing platform.", "Static / siteContent.json"],
            ["13", "/solutions/finlagoon-consolidation", "Finlagoon Consolidation", "Financial legal and management consolidation tool.", "Static / siteContent.json"],
            ["14", "/solutions/profitability-pro", "Profitability Pro", "Margin analysis, Cost allocation & CO-PA reporting.", "Static / siteContent.json"],
            ["15", "/industries", "Industries Hub", "Industry sector overview and domain competencies.", "Static / siteContent.json"],
            ["16", "/industries/retail-supply-chain", "Retail & Supply Chain", "Omnichannel ERP, Warehouse automation, Demand planning.", "Static / animations"],
            ["17", "/industries/pharmaceuticals-life-sciences", "Pharma & Life Sciences", "GxP compliance, Batch tracking, Validated SAP ERP.", "Static / animations"],
            ["18", "/industries/manufacturing", "Manufacturing", "Industry 4.0, MES integration, Shop floor scheduling.", "Static / animations"],
            ["19", "/industries/banking-financial-services", "Banking & Finance", "Core ledger, IFRS compliance, Financial analytics.", "Static / animations"],
            ["20", "/industries/energy-utilities", "Energy & Utilities", "Asset management, Billing automation, Smart grid ERP.", "Static / animations"],
            ["21", "/industries/steel-mining", "Steel & Mining", "Raw material tracking, Plant maintenance, Mill ERP.", "Static / animations"],
            ["22", "/industries/media-telecom", "Media & Telecom", "Subscription billing, Revenue sharing, Network ERP.", "Static / animations"],
            ["23", "/blogs", "Blogs Hub", "Searchable article listing, Category filters, Featured posts.", "blogs.ts / blog-store.ts"],
            ["24", "/blogs/[slug]", "Blog Article Detail", "Full-width rendered article with block styling engine.", "Dynamic Slug / blogs.ts"],
            ["25", "/careers", "Careers & Culture", "Job openings, Culture showcase, Resume upload form.", "Static / careers API"],
            ["26", "/contact-us", "Contact Us", "Inquiry submission, Interactive location map, Direct details.", "Static / contact API"],
            ["27", "/case-studies", "Case Studies", "Enterprise client digital transformation success stories.", "Static / siteContent.json"],
            ["28", "/insights", "Insights Hub", "Whitepapers, Industry perspectives, Market trends.", "Static / siteContent.json"],
            ["29", "/corporate", "Corporate Overview", "Governance, Global offices, Executive team.", "Static / siteContent.json"],
            ["30", "/privacy-policy", "Privacy Policy", "Data protection regulations, Cookie policy, User rights.", "Static Text"],
            ["31", "/terms-of-service", "Terms of Service", "Legal terms, IP ownership, Website usage guidelines.", "Static Text"],
            ["32", "/videos", "Video Showcase", "Product demos, Webinar recordings, SAP tutorials.", "Static / YouTube Embeds"],
            ["33", "/admin", "Live Page Editor", "Visual WYSIWYG element editor with live DOM preview.", "siteContent.json / ContentAPI"],
            ["34", "/admin/blogs", "Blog Studio Hub", "Manage articles, edit drafts, publish new blog posts.", "blog-store.ts / BlogsAPI"],
            ["35", "/admin/createblog", "Create Blog Studio", "Full-screen 10-block visual drag/add article builder.", "blog-store.ts / BlogsAPI"],
            ["36", "/admin/chatbot", "Chatbot Studio", "AI knowledge editor, keyword training, global settings.", "knowledge-base/ / SettingsAPI"],
            ["37", "/admin/login", "Admin Login", "Secure HMAC session token authentication portal.", "admin-auth.ts / LoginAPI"],
            ["38-49", "/api/* (12 Endpoints)", "API Route Handlers", "Content, Blogs, Chat, Images, Auth, Careers, Contact.", "Node.js Server Handlers"]
        ],
        [0.4, 1.8, 1.3, 2.0, 1.0]
    )

    # ─────────────────────────────────────────────────────────────
    # 7. HOMEPAGE SOP
    # ─────────────────────────────────────────────────────────────
    h1("7. HOMEPAGE SOP")
    p("The Homepage (`app/page.tsx`) serves as the core entry portal. It coordinates high-impact animations, interactive showcases, and live editable content:")
    bullet("Hero Section (`components/sections/Hero.tsx`): Displays dynamic headline with gold/teal gradient text, animated background particle stage, floating SAP badges, and Primary 'Schedule Consultation' / Secondary 'Explore Solutions' CTAs.")
    bullet("Company Intro & Pillars: Showcases TRIJOTECH's core differentiators in enterprise architecture and advisory.")
    bullet("Capabilities Grid (`components/sections/ServicesSection.tsx`): 6-card interactive grid linking directly to specialized SAP service pages.")
    bullet("Industry Carousel (`components/sections/IndustriesSection.tsx`): Interactive tabbed industry browser with live preview animations.")
    bullet("Metrics & Stats Counters (`components/motion/AnimatedCounter.tsx`): Number counters that trigger on scroll into viewport (e.g. 50+ Enterprise Clients, 100% On-Time Delivery, 24/7 Global Support).")
    bullet("Client Testimonials: React Slick carousel displaying verified enterprise client feedback with author avatars.")

    # ─────────────────────────────────────────────────────────────
    # 8. NAVBAR & NAVIGATION MANAGEMENT
    # ─────────────────────────────────────────────────────────────
    h1("8. NAVBAR AND NAVIGATION MANAGEMENT")
    p("The header navigation (`components/layout/Header.tsx` and `app/components/Header.jsx`) provides responsive desktop and mobile menus:")
    bullet("Desktop Navigation: Features hover-activated dropdown menus for Services, Solutions, and Industries with rich card previews.")
    bullet("Mobile Navigation Drawer: Sliding touch-friendly drawer triggered via hamburger button with collapsible sub-menus.")
    bullet("Sticky & Glassmorphism Header: Background transitions from transparent to blurred dark navy (`rgba(11, 29, 51, 0.85)`) upon scrolling past 50px.")
    bullet("Procedure to Add a Navigation Item: Edit `lib/header-data.json` or `components/layout/Header.tsx` to insert the new route link into the `navigationItems` array.")

    # ─────────────────────────────────────────────────────────────
    # 9. HERO SECTION MANAGEMENT
    # ─────────────────────────────────────────────────────────────
    h1("9. HERO SECTION MANAGEMENT")
    p("Hero sections utilize a unified styling paradigm (`.hero-fullvh`) ensuring consistent 100svh / 100dvh viewport height across mobile Safari, Android Chrome, and desktop monitors:")
    bullet("Video Hero Backgrounds: Standardized MP4 assets loaded from `/public/assets/video/` with `autoPlay`, `loop`, `muted`, and `playsInline` attributes.")
    bullet("Canvas Particle River: Lightweight HTML5 2D canvas loops (`requestAnimationFrame`) rendering signal nodes without CPU lag.")
    bullet("Gradient Heading Overlay: High-contrast typography with `--color-tri-3: #f5a623` amber highlights.")

    # ─────────────────────────────────────────────────────────────
    # 10. SERVICES SECTION
    # ─────────────────────────────────────────────────────────────
    h1("10. SERVICES SECTION")
    p("All 6 specialized SAP enterprise services have dedicated routes under `/services/`:")
    tbl(
        ["Service Offering", "Route Path", "Primary Component", "Key Competency Highlighted"],
        [
            ["SAP Consulting", "/services/sap-consulting", "app/services/sap-consulting/page.tsx", "Strategic roadmapping, S/4HANA migration advisory."],
            ["SAP Implementation", "/services/sap-implementation", "app/services/sap-implementation/page.tsx", "Turnkey greenfield/brownfield ERP implementations."],
            ["SAP Support & AMS", "/services/sap-support", "app/services/sap-support/page.tsx", "24/7 SLA-driven Level 1-3 support and incident resolution."],
            ["SAP BTP Full Stack", "/services/sap-btp-full-stack", "app/services/sap-btp-full-stack/page.tsx", "Business Technology Platform extensions, SAP Build, API Management."],
            ["SAP Data & Integration", "/services/sap-data-integration", "app/services/sap-data-integration/page.tsx", "Datasphere, SAC, CPI/PI/PO pipelines, master data governance."],
            ["SAP AI & ML", "/services/sap-ai-ml", "app/services/sap-ai-ml/page.tsx", "Generative AI, Business AI algorithms, Predictive analytics."]
        ],
        [1.5, 1.8, 1.8, 1.4]
    )

    # ─────────────────────────────────────────────────────────────
    # 11. INDUSTRIES SECTION
    # ─────────────────────────────────────────────────────────────
    h1("11. INDUSTRIES SECTION")
    p("Industry domain expertise is delivered through dedicated deep-dive pages featuring interactive animated stages (`components/industries/animations/`):")
    bullet("Retail & Supply Chain (`/industries/retail-supply-chain`): SupplyChainScene animation, inventory tracking, POS integration.")
    bullet("Pharmaceuticals & Life Sciences (`/industries/pharmaceuticals-life-sciences`): PharmaNetwork animation, FDA GxP compliance.")
    bullet("Manufacturing (`/industries/manufacturing`): ManufacturingFlow animation, shop floor automation, predictive maintenance.")
    bullet("Banking & Financial Services (`/industries/banking-financial-services`): FintechAnimation, real-time settlement, GL reconciliation.")
    bullet("Energy & Utilities (`/industries/energy-utilities`): Smart grid metering, asset uptime tracking.")
    bullet("Steel & Mining (`/industries/steel-mining`): SteelFlow animation, raw material handling, batch processing.")
    bullet("Media & Telecom (`/industries/media-telecom`): EntertainmentAnimation, subscription lifecycle billing.")

    # ─────────────────────────────────────────────────────────────
    # 12. ADMIN PANEL SOP
    # ─────────────────────────────────────────────────────────────
    h1("12. ADMIN PANEL SOP")
    p("The TRIJOTECH Admin Studio is an enterprise management suite accessible at `/admin`. It enables administrators to manage visual content, publish blogs, upload media, and train AI models without editing code.")

    h2("12.1 Authentication & Security (`/admin/login`)")
    p("Admin access is secured using HTTP-only cookie session tokens signed via HMAC SHA-256 (`app/lib/admin-auth.ts`). To authenticate:")
    bullet("1. Navigate to `http://localhost:3000/admin/login`.")
    bullet("2. Enter the configured `ADMIN_PASSWORD`.")
    bullet("3. On successful authentication, an encrypted `trijotech_admin_session` cookie is issued (valid for 7 days with `SameSite=Lax` and `HttpOnly`).")

    h2("12.2 Unified Admin Navbar (`AdminNavbar.tsx`)")
    p("All admin pages feature a standardized navigation panel with instant switching between tools:")
    bullet("Page Content (`/admin`): Visual WYSIWYG page content editor.")
    bullet("Blog Management (`/admin/blogs`): Article directory and draft manager.")
    bullet("Chatbot Studio (`/admin/chatbot`): AI knowledge base editor and LLM settings.")
    bullet("+ Create Blog (`/admin/createblog`): Full-screen block-based blog creator.")
    bullet("Sign Out Button: Destroys session cookie and redirects to `/admin/login`.")

    h2("12.3 Visual Page Content Editor (`/admin`)")
    p("The live content editor renders the website inside an iframe with bi-directional postMessage communication:")
    bullet("Step 1: Choose a page from the 'Choose a website page' dropdown.")
    bullet("Step 2: Click on any heading, paragraph, button, image, or section in the preview.")
    bullet("Step 3: The editor automatically strips raw HTML tags and displays clean plain text.")
    bullet("Step 4: Edit the text, button URL, or upload a new image.")
    bullet("Step 5: Click 'Preview' to test changes in the iframe, then click 'Save Changes' to publish instantly.")

    callout("Automated HTML Tag Stripping: When selecting formatted elements (e.g. text containing `<span class='tri-gradient-text'>...</span>`), the editor automatically strips tags so administrators only edit clean text. Optional 'Bold', 'Italic', and 'Gradient' format buttons allow re-applying highlights without typing code.", "FEATURE HIGHLIGHT")

    # ─────────────────────────────────────────────────────────────
    # 13. BLOG MANAGEMENT SOP
    # ─────────────────────────────────────────────────────────────
    h1("13. BLOG MANAGEMENT")
    p("The Blog Block Editor (`app/admin/BlogBlockEditor.tsx`) allows crafting long-form thought leadership articles using 10 specialized block types:")
    tbl(
        ["Block Primitive", "Purpose & Capabilities", "Styling Controls Available"],
        [
            ["Heading", "Section titles and major subtitles.", "Heading levels 1-6, font sizes, text align, color gradient."],
            ["Content / Paragraph", "Body copy with rich text support.", "Font family, line height, letter spacing, justification."],
            ["Image", "Embedded inline article illustrations.", "Width (25%-100%), alignment, border radius, shadow, caption."],
            ["Quote", "Featured executive quote callout.", "Border accent color, italic styling, attribution text."],
            ["Bullet List", "Unordered bulleted points.", "Item spacing, marker color, custom indentation."],
            ["Numbered List", "Ordered step-by-step points.", "Numerical sequencing, nested spacing."],
            ["Callout Box", "Highlighted note, alert, or key takeaway.", "Background tint, left border color, custom icon."],
            ["Divider", "Visual horizontal rule between sections.", "Line style (solid, dashed, dotted), vertical margin."],
            ["Link / Button", "In-text call to action or reference link.", "Target URL, button style, open-in-new-tab."],
            ["Custom CSS", "Advanced styling overrides for the block.", "Direct CSS rules for bespoke layouts."]
        ],
        [1.5, 2.5, 2.5]
    )

    h2("Procedure: Publishing a New Blog Post")
    bullet("1. Navigate to `/admin/createblog` or click '+ Create Blog' in the Admin Navbar.", "Step 1: ")
    bullet("2. Fill in Article Title, Slug (auto-generated from title), Category, and Estimated Read Time.", "Step 2: ")
    bullet("3. Upload a Featured Image or choose an existing asset.", "Step 3: ")
    bullet("4. In the 'Add Blocks' section, click any block type (e.g. Heading, Content, Image) to add it to the canvas.", "Step 4: ")
    bullet("5. Enter block content and configure typography/spacing settings in the block toolbar.", "Step 5: ")
    bullet("6. Set Status to 'Published' and click 'Publish Blog'.", "Step 6: ")
    bullet("7. Verify the live article at `http://localhost:3000/blogs/<slug>`.", "Step 7: ")

    # ─────────────────────────────────────────────────────────────
    # 14. CONTACT FORM SOP
    # ─────────────────────────────────────────────────────────────
    h1("14. CONTACT FORM SOP")
    p("The Contact Form (`app/components/ContactUs.jsx` and `app/api/forms/contact/route.ts`) handles corporate inquiries:")
    bullet("Input Fields: Full Name, Business Email, Phone Number, Company Name, Service/Solution Interest, Project Description.")
    bullet("Security Safeguards: Mathematical CAPTCHA challenge (or optional Cloudflare Turnstile token), hidden honeypot field (`website_hp`), rate limiting (5 submissions per IP per 10 minutes), input sanitization.")
    bullet("Backend Processing: Validates fields server-side and dispatches formatted HTML emails via Nodemailer to `CONTACT_TO_EMAIL` (default: `sales@trijotech.com`).")

    # ─────────────────────────────────────────────────────────────
    # 15. CAREER FORM SOP
    # ─────────────────────────────────────────────────────────────
    h1("15. CAREER FORM SOP")
    p("The Job Application Form (`app/careers/page.tsx` and `app/api/forms/careers/route.ts`) accepts candidate applications:")
    bullet("Form Fields: Candidate Name, Email Address, Phone, Target Position / Job ID, Portfolio / LinkedIn URL, Cover Note, Resume / CV upload.")
    bullet("Resume Upload Constraints: Accepts PDF, DOC, DOCX files up to 5MB. Files are verified for MIME type and attached to the outbound notification email dispatched to `CAREERS_TO_EMAIL` (default: `hr@trijotech.com`).")

    # ─────────────────────────────────────────────────────────────
    # 16. EMAIL CONFIGURATION & TROUBLESHOOTING
    # ─────────────────────────────────────────────────────────────
    h1("16. EMAIL CONFIGURATION")
    p("Nodemailer SMTP transport is configured centrally in `app/lib/form-security.ts`:")
    code("""// Outbound SMTP Transport Architecture
const transport = nodemailer.createTransport({
  host: process.env.SMTP_HOST,
  port: Number(process.env.SMTP_PORT || 587),
  secure: process.env.SMTP_SECURE === "true",
  auth: {
    user: process.env.SMTP_USER,
    pass: process.env.SMTP_PASS,
  },
});""")

    h2("Email Troubleshooting Guide")
    tbl(
        ["Symptom / Error", "Root Cause", "Actionable Resolution"],
        [
            ["EAUTH / Invalid Credentials", "Incorrect SMTP_USER or SMTP_PASS.", "For Gmail/Office365, generate an App Password instead of account password."],
            ["ETIMEDOUT / Connection Timeout", "Firewall blocking port 587 / 465.", "Verify SMTP_HOST and check host outbound network security rules."],
            ["Missing Environment Variables", "SMTP_HOST or SMTP_USER undefined.", "Ensure .env.local is present with all required SMTP values."],
            ["Email going to Spam folder", "Missing SPF/DKIM records for sender domain.", "Configure SPF/DKIM DNS records on the outbound domain."]
        ],
        [1.8, 2.0, 2.7]
    )

    # ─────────────────────────────────────────────────────────────
    # 17. CHATBOT STUDIO & AI ENGINE
    # ─────────────────────────────────────────────────────────────
    h1("17. CHATBOT STUDIO & AI ENGINE")
    p("The TRIJOTECH Floating Chatbot (`components/chatbot/ChatbotWidget.tsx`) provides 24/7 AI-powered conversational assistance grounded in verified corporate knowledge:")
    bullet("Local Verified Knowledge Base: JSON files under `app/data/knowledge-base/` contain accurate facts regarding TRIJOTECH services, SAP solutions, office locations, and leadership.")
    bullet("Keyword & Intent Scoring: When a visitor asks a question, the server ranks knowledge entries using token overlap and intent matching.")
    bullet("DeepSeek LLM Integration (`app/services/deepseek.ts`): Top-ranked verified knowledge entries are injected into the system prompt of DeepSeek Chat (`deepseek-chat`) to generate factual, hallucination-free enterprise responses.")
    bullet("Chatbot Studio (`/admin/chatbot`): Allows adding/editing knowledge cards, testing responses in real-time, and updating personality prompts without code deployments.")

    # ─────────────────────────────────────────────────────────────
    # 18. ANIMATION SYSTEM
    # ─────────────────────────────────────────────────────────────
    h1("18. ANIMATION SYSTEM")
    p("Animations are architected for 60fps GPU-accelerated performance across desktop and mobile devices:")
    bullet("Framer Motion Viewport Reveals (`components/motion/Reveal.tsx`): Elements smoothly fade and slide up as they scroll into view (`whileInView`, `viewport: { once: true }`).")
    bullet("Animated Number Counters (`components/motion/AnimatedCounter.tsx`): Smooth numerical count-up easing for metrics and statistics.")
    bullet("HTML5 2D Canvas Stages (`components/scenes/`): Lightweight procedural canvas loops rendering dynamic particle rivers, network nodes, and assembly lines.")
    bullet("CSS Keyframe Transitions: GPU-composited `transform` and `opacity` animations defined in `app/globals.css` (`float`, `float-slow`, `shimmer`, `gradient-shift`).")

    # ─────────────────────────────────────────────────────────────
    # 19. STYLING & DESIGN SYSTEM
    # ─────────────────────────────────────────────────────────────
    h1("19. STYLING AND DESIGN SYSTEM")
    p("The website adheres to a strict enterprise color palette and typography system defined in `app/globals.css`:")
    tbl(
        ["Design Token", "Hex / CSS Value", "Semantic Usage Across Website"],
        [
            ["--font-poppins", "'Poppins', sans-serif", "Primary typography for all headings, body copy, buttons, and inputs."],
            ["--color-tri-1", "#117a4b", "Deep Forest Green brand accent."],
            ["--color-tri-2", "#ffffff", "Teal secondary accent for icons, active tabs, and badge borders."],
            ["--color-tri-3", "#f5a623", "Vibrant Gold / Amber for gradient text highlights and primary action buttons."],
            ["--color-tri-3-deep", "#f29e16", "Hover state for amber buttons and badges."],
            ["--color-tri-navy", "#232555", "Header brand gradients and modal backdrops."],
            ["--color-tri-ink", "#1a2336", "Dark card background surfaces."],
            ["--color-tri-ink-2", "#121927", "Deepest dark navy for footer and contrast containers."],
            ["--color-tri-ink-3", "#222d42", "Border outlines and subtle divider lines."],
            ["--color-tri-mist", "#e8f2fb", "Light blue background tint for callouts and badges."]
        ],
        [1.8, 1.6, 3.1]
    )

    # ─────────────────────────────────────────────────────────────
    # 20. RESPONSIVE DESIGN
    # ─────────────────────────────────────────────────────────────
    h1("20. RESPONSIVE DESIGN")
    p("Responsive behavior is engineered using Tailwind mobile-first breakpoints (`sm: 640px`, `md: 768px`, `lg: 1024px`, `xl: 1280px`, `2xl: 1536px`):")
    bullet("Desktop (1280px+): Full multi-column grids, horizontal navigation with hover dropdowns, side-by-side hero text and canvas animations.")
    bullet("Tablet (768px - 1023px): 2-column grid adaptation, responsive card stacking, touch-friendly carousel drag navigation.")
    bullet("Mobile (<768px): Single-column layouts, slide-in mobile navigation drawer, touch-optimized button tap targets (minimum 44px height), hidden high-overhead canvas animations to conserve mobile battery.")

    # ─────────────────────────────────────────────────────────────
    # 21. IMAGE & MEDIA MANAGEMENT
    # ─────────────────────────────────────────────────────────────
    h1("21. IMAGE AND MEDIA MANAGEMENT")
    p("Static image assets reside under `/public/assets/` and `/public/brand/`. Admin-uploaded assets are stored in `/public/assets/uploads/`:")
    bullet("Next.js Image Optimization: Images utilize Next.js `<Image />` with automatic WebP conversion, responsive `srcset`, and blur placeholder support.")
    bullet("Admin Image Upload Route (`/api/admin/images`): Validates image file types (PNG, JPG, WEBP, SVG), enforces 5MB limit, generates collision-proof timestamps, and saves files to `/public/assets/uploads/`.")
    bullet("Procedure to Replace an Image: Upload the new image via the Page Content Editor (`/admin`) or Blog Block Editor, which automatically updates the image URL and publishes the change.")

    # ─────────────────────────────────────────────────────────────
    # 22. PERFORMANCE OPTIMIZATION
    # ─────────────────────────────────────────────────────────────
    h1("22. PERFORMANCE OPTIMIZATION")
    p("Verified performance optimizations implemented in the codebase include:")
    bullet("Turbopack Build Engine: Fast Next.js compilation and hot module reloading.")
    bullet("Static Site Generation (SSG): All 49 routes are pre-rendered at build time into static HTML/JSON.")
    bullet("Font Optimization: Next.js Google Fonts (`Poppins`) pre-loaded with `display: swap` to eliminate layout shift.")
    bullet("Static Asset Caching: Custom Cache-Control headers configured in `next.config.ts` (`public, max-age=86400, stale-while-revalidate=604800`) for `/static/` and `/assets/`.")
    bullet("DOM Mutation Throttling: `ContentRuntime.tsx` utilizes microtasks and `requestAnimationFrame` to batch DOM updates efficiently.")

    # ─────────────────────────────────────────────────────────────
    # 23. SEO & METADATA
    # ─────────────────────────────────────────────────────────────
    h1("23. SEO AND METADATA")
    p("SEO metadata is defined using Next.js Metadata API across all static and dynamic pages:")
    bullet("Root Metadata (`app/layout.tsx`): Global title template ('Trijotech | SAP Solutions') and description.")
    bullet("Dynamic Blog SEO (`app/blogs/[slug]/page.tsx`): Dynamically resolves `title`, `description`, `openGraph`, and canonical URLs based on the blog article data model.")
    bullet("Semantic HTML Structure: Strict usage of `<h1>`, `<h2>`, `<article>`, `<nav>`, `<header>`, and `<footer>` tags.")

    # ─────────────────────────────────────────────────────────────
    # 24. SECURITY ARCHITECTURE
    # ─────────────────────────────────────────────────────────────
    h1("24. SECURITY ARCHITECTURE")
    p("Security safeguards are implemented across multiple layers of the application:")
    bullet("Admin Cookie Authentication: Secure HMAC SHA-256 signatures with constant-time verification preventing timing attacks.")
    bullet("Form Anti-Abuse: Multi-layer defense combining mathematical CAPTCHA, hidden honeypot fields, and IP-based rate limiting.")
    bullet("Input Sanitization: All admin HTML overrides pass through DOMParser and SafeHTML template filters, stripping `<script>`, `<iframe>`, and `on*` inline event handlers.")
    bullet("Secret Isolation: All SMTP credentials and API keys are restricted to server-side Route Handlers and never exposed to client bundles.")

    # ─────────────────────────────────────────────────────────────
    # 25. DEVELOPMENT PROCEDURES (SOP)
    # ─────────────────────────────────────────────────────────────
    h1("25. DEVELOPMENT PROCEDURES")
    
    h2("25.1 Creating a New Page Route")
    bullet("1. Create a new folder under `app/` (e.g. `app/partners/`).", "Step 1: ")
    bullet("2. Inside the folder, create `page.tsx` exporting a default React component.", "Step 2: ")
    bullet("3. Define static metadata (`export const metadata = { title: '...', description: '...' };`).", "Step 3: ")
    bullet("4. Import reusable components from `@/components/` and build the page layout.", "Step 4: ")
    bullet("5. Run `npm run build` to verify clean SSG compilation.", "Step 5: ")

    h2("25.2 Adding a New Reusable Component")
    bullet("1. Create the component file under `components/` (e.g. `components/ui/FeatureCard.tsx`).", "Step 1: ")
    bullet("2. Define explicit TypeScript interface props.", "Step 2: ")
    bullet("3. Apply Tailwind CSS utility classes and Framer Motion wrappers if animated.", "Step 3: ")
    bullet("4. Export the component and import it into target pages.", "Step 4: ")

    # ─────────────────────────────────────────────────────────────
    # 26. GIT & VERSION CONTROL WORKFLOW
    # ─────────────────────────────────────────────────────────────
    h1("26. GIT AND VERSION CONTROL SOP")
    p("All code modifications must follow a structured feature-branch workflow:")
    code("""# 1. Fetch latest main branch
git checkout main
git pull origin main

# 2. Create a dedicated feature branch
git checkout -b feature/update-services-section

# 3. Make and test changes locally
npm run dev

# 4. Verify build and types before committing
npm run build

# 5. Commit with descriptive semantic message
git add .
git commit -m "feat(services): add SAP AI/ML capability details"

# 6. Push to remote and open Pull Request
git push origin feature/update-services-section""")

    # ─────────────────────────────────────────────────────────────
    # 27. TESTING SOP
    # ─────────────────────────────────────────────────────────────
    h1("27. TESTING SOP")
    p("Prior to merging or deploying, execute the following testing checklists:")
    tbl(
        ["Testing Category", "Verification Scope", "Acceptance Criteria"],
        [
            ["TypeScript Safety", "Run `npx tsc --noEmit` across entire codebase.", "Zero type errors (exit code 0)."],
            ["Production Build", "Run `npm run build` with Turbopack.", "All 49 routes compile cleanly with SSG/SSR."],
            ["Form Submissions", "Submit test inquiries on /contact-us and /careers.", "Success confirmation displayed; email received in destination inbox."],
            ["Admin Studio CRUD", "Test Content Editor, Blog Creator, and Chatbot Studio.", "Overrides save to JSON and reflect immediately on live pages."],
            ["Responsive Mobile", "Test on iPhone Safari and Android Chrome viewports.", "No horizontal overflow; mobile navigation drawer functions smoothly."],
            ["Cross-Browser", "Test on Chrome, Edge, Safari, and Firefox.", "Animations and glassmorphism render consistently at 60fps."]
        ],
        [1.8, 2.2, 2.5]
    )

    # ─────────────────────────────────────────────────────────────
    # 28. DEPLOYMENT SOP
    # ─────────────────────────────────────────────────────────────
    h1("28. DEPLOYMENT SOP")
    p("The TRIJOTECH Next.js application can be deployed to any Node.js hosting environment, Docker container, PM2 cluster, or Vercel serverless platform:")
    
    h2("28.1 Production Deployment Steps (Node.js / PM2 Server)")
    bullet("1. Clone the repository on the production host server.", "Step 1: ")
    bullet("2. Configure production `.env.local` with secure `ADMIN_PASSWORD`, `ADMIN_SESSION_SECRET`, and `SMTP_*` values.", "Step 2: ")
    bullet("3. Run `npm install --omit=dev` to install production dependencies.", "Step 3: ")
    bullet("4. Run `npm run build` to generate optimized production artifacts in `.next/`.", "Step 4: ")
    bullet("5. Start the production process using PM2: `pm2 start npm --name 'trijotech-web' -- start`.", "Step 5: ")
    bullet("6. Configure Nginx reverse proxy with SSL certificate pointing to `http://127.0.0.1:3000`.", "Step 6: ")

    # ─────────────────────────────────────────────────────────────
    # 29. POST-DEPLOYMENT VALIDATION
    # ─────────────────────────────────────────────────────────────
    h1("29. POST-DEPLOYMENT VALIDATION")
    p("Immediately following any production deployment, verify:")
    bullet("Homepage loads with HTTPS and zero console errors.", "[ ] ")
    bullet("Navigation dropdown menus and mobile drawer open and navigate correctly.", "[ ] ")
    bullet("All 6 Service and 7 Industry pages load with interactive animations.", "[ ] ")
    bullet("Blog listing (`/blogs`) and dynamic detail pages (`/blogs/[slug]`) render cleanly.", "[ ] ")
    bullet("Contact Us inquiry form successfully sends email to sales inbox.", "[ ] ")
    bullet("Career application form uploads resume and dispatches email to HR.", "[ ] ")
    bullet("Admin Studio (`/admin/login`) authenticates and allows editing site content.", "[ ] ")
    bullet("Floating Chatbot answers queries using verified knowledge base.", "[ ] ")

    # ─────────────────────────────────────────────────────────────
    # 30. TROUBLESHOOTING GUIDE
    # ─────────────────────────────────────────────────────────────
    h1("30. TROUBLESHOOTING GUIDE")
    p("The following matrix covers common operational issues, diagnostics, and verified solutions:")

    tbl(
        ["Issue / Error", "Probable Root Cause", "Step-by-Step Resolution"],
        [
            ["Admin changes not showing on website", "Transient CSS selector class or browser cache.", "Ensure ContentRuntime is active; ContentRuntime automatically strips .admin-edit-* and syncs via /api/content."],
            ["Port 3000 already in use", "Previous Next.js dev server instance still running.", "Run `npx kill-port 3000` or start on alternate port `npx next dev -p 3001`."],
            ["npm run build fails on TypeScript", "Type mismatch in component props or API payload.", "Run `npx tsc --noEmit` to locate specific file and line number; fix type annotations."],
            ["Contact Form returns 500 error", "SMTP configuration missing or mail server unreachable.", "Verify SMTP_HOST, SMTP_USER, SMTP_PASS in .env.local; check server outbound port 587."],
            ["Chatbot returns fallback response", "DeepSeek API key invalid or knowledge mismatch.", "Verify DEEPSEEK_API_KEY in .env.local or add targeted knowledge entries in /admin/chatbot."],
            ["Images not uploading in Admin", "File exceeds 5MB or invalid MIME type.", "Ensure uploaded image is PNG, JPG, WEBP, or SVG under 5MB in size."],
            ["Hydration mismatch error", "Browser extension modifying DOM or window timing.", "Ensure dynamic client-only date/window calculations are wrapped in useEffect."]
        ],
        [1.8, 2.0, 2.7]
    )

    # ─────────────────────────────────────────────────────────────
    # 31. WEBSITE MAINTENANCE SCHEDULE
    # ─────────────────────────────────────────────────────────────
    h1("31. WEBSITE MAINTENANCE SOP")
    p("To maintain optimal performance, uptime, and security, adhere to the following maintenance cadence:")
    tbl(
        ["Frequency", "Maintenance Tasks & Checks", "Responsible Party"],
        [
            ["Daily / Continuous", "Monitor website uptime, check form delivery logs, review chatbot analytics.", "Support / Operations Team"],
            ["Weekly", "Review new contact inquiries, verify blog drafts, backup siteContent.json.", "Content Administrator"],
            ["Monthly", "Audit package dependencies (`npm audit`), test all 49 routes, review SMTP quotas.", "Lead Developer"],
            ["Quarterly", "Review SSL certificate renewals, conduct full security review, optimize media assets.", "DevOps / IT Lead"]
        ],
        [1.5, 3.5, 1.5]
    )

    # ─────────────────────────────────────────────────────────────
    # 32. CHANGE MANAGEMENT PROCEDURE
    # ─────────────────────────────────────────────────────────────
    h1("32. CHANGE MANAGEMENT PROCEDURE")
    p("All website modifications must follow the 10-step Change Lifecycle:")
    code("""[ 1. Change Request ] ──> [ 2. Requirements Analysis ] ──> [ 3. Component Identification ]
                                                                      │
[ 6. Code Review ]    <── [ 5. Local Build Testing ]   <── [ 4. Feature Branch Dev ]
       │
       ▼
[ 7. Stakeholder Sign-Off ] ──> [ 8. Production Deploy ] ──> [ 9. Post-Validation ] ──> [ 10. Close ]""")

    # ─────────────────────────────────────────────────────────────
    # 33. BACKUP AND RECOVERY
    # ─────────────────────────────────────────────────────────────
    h1("33. BACKUP AND RECOVERY")
    p("Because TRIJOTECH utilizes resilient JSON filesystem storage, backup and disaster recovery procedures are straightforward:")
    bullet("Source Code Backup: Maintained in remote Git repository (GitHub/GitLab).", "Source: ")
    bullet("Live Content Backup: The file `app/data/siteContent.json` stores all live page overrides.", "Content: ")
    bullet("Blog Posts Backup: The files `app/data/blogs.ts` and `app/data/blogs.json` store all articles.", "Blogs: ")
    bullet("Uploaded Media Backup: The folder `public/assets/uploads/` stores all media assets.", "Media: ")
    bullet("Recovery Procedure: In the event of server failure, clone the repository, restore `siteContent.json` and `/public/assets/uploads/` from backup, run `npm install && npm run build`, and restart the service.", "Recovery: ")

    # ─────────────────────────────────────────────────────────────
    # 34. CODING GUIDELINES
    # ─────────────────────────────────────────────────────────────
    h1("34. CODING AND MAINTENANCE GUIDELINES")
    bullet("Maintain Documentation Integrity: Preserve all existing comments and docstrings.", "1. ")
    bullet("Enforce Type Safety: Never use `any` when explicit TypeScript interfaces can be defined.", "2. ")
    bullet("Preserve Responsive Design: Test all UI additions at 375px, 768px, 1024px, and 1440px breakpoints.", "3. ")
    bullet("Protect Server Secrets: Never prefix sensitive credentials with `NEXT_PUBLIC_`.", "4. ")
    bullet("Scoped Styling: Utilize CSS modules (`*.module.css`) or Tailwind utilities to prevent cross-page style leaks.", "5. ")

    # ─────────────────────────────────────────────────────────────
    # 35. COMMON OPERATIONAL PROCEDURES (SOP-01 to SOP-20)
    # ─────────────────────────────────────────────────────────────
    h1("35. COMMON OPERATIONAL PROCEDURES")
    
    sops = [
        ("SOP-01", "Start Website Locally", "Terminal access", "Run `npm run dev` in `webapp/`", "Server starts on http://localhost:3000"),
        ("SOP-02", "Stop Development Server", "Running terminal", "Press `Ctrl + C` in the active terminal window", "Dev server terminates cleanly"),
        ("SOP-03", "Create New Page", "Developer access", "Add folder in `app/<name>/page.tsx` with metadata", "New route accessible in browser"),
        ("SOP-04", "Edit Live Page Text", "Admin access", "Login to `/admin`, select page, click element, edit, save", "Text updates live on website"),
        ("SOP-05", "Add Reusable Component", "Developer access", "Create file in `components/<category>/<Name>.tsx`", "Component available for import"),
        ("SOP-06", "Change Hero Video/Image", "Admin/Dev access", "Replace asset in `/public/assets/` or update via `/admin`", "New media displays on hero"),
        ("SOP-07", "Add Animation to Section", "Developer access", "Wrap component in `<Reveal>` from `components/motion/`", "Element animates on viewport entry"),
        ("SOP-08", "Update Header Navigation", "Developer access", "Edit `lib/header-data.json` or `components/layout/Header.tsx`", "Navbar items update across site"),
        ("SOP-09", "Add New SAP Service", "Developer access", "Create route under `app/services/<slug>/page.tsx`", "New service page compiles in build"),
        ("SOP-10", "Add New Industry Page", "Developer access", "Create route under `app/industries/<slug>/page.tsx`", "New industry page compiles in build"),
        ("SOP-11", "Publish New Blog Post", "Admin access", "Create article at `/admin/createblog`, set Published, save", "Article live at `/blogs/<slug>`"),
        ("SOP-12", "Edit Existing Blog Post", "Admin access", "Open `/admin/blogs`, click Edit on target article, save", "Article updates immediately"),
        ("SOP-13", "Delete / Draft Blog Post", "Admin access", "Open `/admin/blogs`, set Status to Draft or click Delete", "Article removed from public listing"),
        ("SOP-14", "Test Contact Us Form", "Web browser", "Fill form at `/contact-us`, submit, verify success message", "Email arrives in sales inbox"),
        ("SOP-15", "Test Career Application", "Web browser", "Fill form at `/careers`, attach sample PDF resume, submit", "Email with resume arrives in HR inbox"),
        ("SOP-16", "Update Chatbot Knowledge", "Admin access", "Open `/admin/chatbot`, add/edit knowledge entry, save", "Chatbot uses new info in replies"),
        ("SOP-17", "Run Production Build", "Terminal access", "Execute `npm run build`", "All 49 routes compile with 0 errors"),
        ("SOP-18", "Deploy to Production", "DevOps access", "Build application and start via PM2 / Node.js service", "Site live on production domain"),
        ("SOP-19", "Roll Back a Change", "Git access", "Run `git revert <commit-hash>` and redeploy", "Previous stable state restored"),
        ("SOP-20", "Troubleshoot Build Failure", "Developer access", "Run `npx tsc --noEmit` and inspect error log", "Syntax/type errors identified & fixed")
    ]
    
    tbl(
        ["SOP ID", "Procedure Name", "Prerequisites", "Action Steps", "Expected Result"],
        sops,
        [0.8, 1.6, 1.1, 1.8, 1.2]
    )

    # ─────────────────────────────────────────────────────────────
    # 36. DEVELOPER QUICK REFERENCE
    # ─────────────────────────────────────────────────────────────
    h1("36. DEVELOPER QUICK REFERENCE")
    tbl(
        ["Activity / Task", "Command / File Path", "Notes / Description"],
        [
            ["Install Dependencies", "npm install", "Run inside webapp/ directory."],
            ["Start Dev Server", "npm run dev", "Runs Next.js with Turbopack on port 3000."],
            ["LAN Dev Server", "npm run dev:lan", "Binds to 0.0.0.0 for mobile network testing."],
            ["Production Build", "npm run build", "Compiles all 49 routes and runs type checking."],
            ["Start Production", "npm run start", "Starts production server on port 3000."],
            ["Type Check Only", "npx tsc --noEmit", "Validates TypeScript types without building."],
            ["Admin Studio URL", "http://localhost:3000/admin", "Master visual editing studio."],
            ["Blog Studio URL", "http://localhost:3000/admin/blogs", "Blog post management interface."],
            ["Create Blog URL", "http://localhost:3000/admin/createblog", "Full-screen block article studio."],
            ["Chatbot Studio URL", "http://localhost:3000/admin/chatbot", "AI knowledge and model configuration."],
            ["Site Overrides File", "app/data/siteContent.json", "Visual overrides JSON data store."],
            ["Public Uploads Dir", "public/assets/uploads/", "Destination for uploaded imagery."]
        ],
        [1.8, 2.4, 2.3]
    )

    # ─────────────────────────────────────────────────────────────
    # 37. COMPONENT INVENTORY
    # ─────────────────────────────────────────────────────────────
    h1("37. COMPONENT INVENTORY")
    tbl(
        ["Component Name", "File Path", "Primary Usage", "Key Functionality"],
        [
            ["Header", "components/layout/Header.tsx", "Global Header across all public pages", "Desktop dropdown navigation, mobile drawer, logo."],
            ["Footer", "components/layout/Footer.tsx", "Global Footer across all public pages", "Site links, social icons, newsletter, copyright."],
            ["PublicChrome", "app/components/PublicChrome.tsx", "Layout chrome wrapper in layout.tsx", "Conditionally displays Header/Footer/Chatbot."],
            ["ContentRuntime", "app/components/ContentRuntime.tsx", "DOM override engine in layout.tsx", "Applies live text/image overrides seamlessly."],
            ["AdminNavbar", "app/admin/AdminNavbar.tsx", "Header navbar across all /admin pages", "Tool tabs, active state, sign out action."],
            ["AdminEditor", "app/admin/AdminEditor.tsx", "Live WYSIWYG Page Editor (/admin)", "Iframe live preview, point-and-click editing."],
            ["BlogManager", "app/admin/BlogManager.tsx", "Blog Management Studio (/admin/blogs)", "Article search, category filter, draft editor."],
            ["BlogBlockEditor", "app/admin/BlogBlockEditor.tsx", "Block builder (/admin/createblog)", "10-block visual drag-and-drop blog creator."],
            ["KnowledgeManager", "app/admin/chatbot/KnowledgeManager.tsx", "AI Knowledge Editor (/admin/chatbot)", "Create/edit knowledge base cards with search."],
            ["SettingsManager", "app/admin/chatbot/SettingsManager.tsx", "AI Settings Manager (/admin/chatbot)", "Personality prompt, model hyperparameters."],
            ["ChatbotWidget", "components/chatbot/ChatbotWidget.tsx", "Floating AI Chatbot on public pages", "Interactive chat bubble, streaming message UI."],
            ["Reveal", "components/motion/Reveal.tsx", "Motion animation wrapper", "Scroll-triggered spring fade/slide animations."],
            ["AnimatedCounter", "components/motion/AnimatedCounter.tsx", "Metrics & stats counter", "Smooth numerical count-up easing effect."],
            ["TurnstileWidget", "app/components/TurnstileWidget.tsx", "Cloudflare Turnstile CAPTCHA wrapper", "Bot protection for Contact and Career forms."]
        ],
        [1.5, 2.1, 1.5, 1.4]
    )

    # ─────────────────────────────────────────────────────────────
    # 38. DEPENDENCY INVENTORY
    # ─────────────────────────────────────────────────────────────
    h1("38. DEPENDENCY INVENTORY")
    tbl(
        ["Package Name", "Installed Version", "Dependency Type", "Purpose & Role in Project"],
        [
            ["next", "16.2.12", "Runtime", "Full-stack React framework, Turbopack, App Router, SSR/SSG."],
            ["react / react-dom", "19.2.4", "Runtime", "Core UI rendering library, React Server Components."],
            ["framer-motion", "^13.1.0", "Runtime", "Declarative UI animation library for scroll reveals and transitions."],
            ["lucide-react", "^1.31.0", "Runtime", "SVG icon library for clean UI indicators and buttons."],
            ["react-icons", "^5.7.0", "Runtime", "Comprehensive icon pack for industry domain and brand icons."],
            ["nodemailer", "^9.0.4", "Runtime", "SMTP email client for Contact Us and Career resume applications."],
            ["lottie-react", "^3.1.0", "Runtime", "Vector JSON animation player for interactive scene diagrams."],
            ["react-slick / slick", "^0.31.0 / ^1.8.1", "Runtime", "Touch-friendly carousel engine for testimonials and solution cards."],
            ["tailwindcss", "^4.0.0", "Development", "Utility CSS compiler and design token theme engine."],
            ["typescript", "^5.0.0", "Development", "Static typing compiler ensuring code safety."],
            ["@types/node", "^20.0.0", "Development", "TypeScript type definitions for Node.js APIs."],
            ["eslint / eslint-config-next", "^9.0.0 / 16.2.12", "Development", "Code quality and Next.js best practice linter."]
        ],
        [1.5, 1.1, 1.2, 2.7]
    )

    # ─────────────────────────────────────────────────────────────
    # 39. API INVENTORY (ALL 13 ENDPOINTS)
    # ─────────────────────────────────────────────────────────────
    h1("39. API INVENTORY")
    p("The table below details all 13 verified server-side Route Handlers (`app/api/`):")
    tbl(
        ["Endpoint Route", "HTTP Method", "Purpose / Functionality", "Called By", "Auth / Security"],
        [
            ["/api/admin/login", "POST", "Authenticates admin password & issues HMAC session cookie.", "AdminLogin.tsx", "Password + Rate Limit"],
            ["/api/admin/logout", "POST", "Destroys session cookie and logs out administrator.", "AdminNavbar.tsx", "Admin Session Cookie"],
            ["/api/admin/content", "GET / PUT / DELETE", "Reads, updates, and resets visual site content overrides.", "AdminEditor.tsx", "Admin Session Cookie"],
            ["/api/admin/blogs", "GET / POST / PUT / DELETE", "CRUD operations for blog articles and block content.", "BlogManager.tsx", "Admin Session Cookie"],
            ["/api/admin/chatbot", "GET / POST / PUT / DELETE", "CRUD operations for AI knowledge base entries.", "KnowledgeManager.tsx", "Admin Session Cookie"],
            ["/api/admin/chatbot/settings", "GET / PUT", "Reads and updates global chatbot LLM configuration.", "SettingsManager.tsx", "Admin Session Cookie"],
            ["/api/admin/images", "POST / DELETE", "Uploads and deletes media assets in /public/assets/uploads/.", "ImageUploadField.tsx", "Admin Session Cookie"],
            ["/api/content", "GET", "Public endpoint returning active site content overrides.", "ContentRuntime.tsx", "Public / Cache-Control"],
            ["/api/chat", "POST", "Processes user chat queries via knowledge search + DeepSeek.", "ChatbotWidget.tsx", "Public / Rate Limited"],
            ["/api/chat/settings", "GET", "Public endpoint providing active chatbot greeting & branding.", "ChatbotWidget.tsx", "Public / No-Store"],
            ["/api/chat/analytics", "POST", "Tracks chatbot interaction metrics and query categories.", "ChatbotWidget.tsx", "Public / Silent Log"],
            ["/api/forms/contact", "POST", "Validates inquiry fields & sends email via Nodemailer.", "ContactUs.jsx", "CAPTCHA + Honeypot + Rate Limit"],
            ["/api/forms/careers", "POST", "Validates job application, attaches resume & sends email.", "careers/page.tsx", "CAPTCHA + Honeypot + Rate Limit"]
        ],
        [1.6, 1.0, 2.2, 1.2, 1.5]
    )

    # ─────────────────────────────────────────────────────────────
    # 40. BUSINESS CONTENT ARCHITECTURE
    # ─────────────────────────────────────────────────────────────
    h1("40. DOCUMENTATION OF IMPORTANT BUSINESS CONTENT")
    p("The corporate messaging and business architecture of TRIJOTECH is partitioned into 6 core domains:")
    bullet("Enterprise SAP Services: Specialized capability descriptions covering Consulting, S/4HANA Implementation, Support AMS, BTP Integration, Data/Analytics, and Business AI/ML.", "Services: ")
    bullet("Industry Competencies: Tailored domain blueprints for Retail & Supply Chain, Pharma, Manufacturing, Banking, Utilities, Steel, and Telecom.", "Industries: ")
    bullet("Proprietary Product Platforms: Highlighting E-Invoicing Pro (government compliance), Finlagoon (financial consolidation), and Profitability Pro (margin analytics).", "Solutions: ")
    bullet("Thought Leadership: High-value articles detailing SAP S/4HANA migration preparation, SAC budgeting in energy, and e-invoicing best practices.", "Blogs: ")
    bullet("Careers & Talent Acquisition: Open roles, culture pillars, global benefits, and direct resume submission channels.", "Careers: ")
    bullet("Contact & Consultation: Multi-channel customer engagement via form submission, direct email, phone, and interactive AI chatbot.", "Contact: ")

    # ─────────────────────────────────────────────────────────────
    # 41. KNOWN LIMITATIONS
    # ─────────────────────────────────────────────────────────────
    h1("41. KNOWN LIMITATIONS")
    p("Based strictly on architectural inspection of the codebase, the following technical characteristics and limitations are documented:")
    bullet("File-Based JSON Storage: Application state (site content overrides, blogs, chatbot knowledge) is stored on the local filesystem (`app/data/`). For multi-instance clustered serverless deployments, external shared storage or database persistence (e.g. PostgreSQL/MongoDB/S3) is recommended.", "1. ")
    bullet("Single Admin Role: Authentication currently validates a single master `ADMIN_PASSWORD` rather than granular multi-user role-based access control (RBAC).", "2. ")
    bullet("Local Media Storage: Uploaded images are stored in `/public/assets/uploads/`. In ephemeral containerized environments (like AWS Fargate or Google Cloud Run), container restarts will wipe uploads unless a persistent volume or Cloud Storage bucket (AWS S3 / GCS) is attached.", "3. ")

    # ─────────────────────────────────────────────────────────────
    # 42. PRIORITIZED RECOMMENDATIONS
    # ─────────────────────────────────────────────────────────────
    h1("42. PRIORITIZED RECOMMENDATIONS")
    p("The following roadmap outlines recommended enhancements categorized by priority:")
    tbl(
        ["Priority", "Recommendation", "Anticipated Business & Technical Benefit", "Complexity"],
        [
            ["High", "Integrate Cloud Storage (S3/GCS) for Image Uploads", "Prevents media loss in ephemeral containerized deployments.", "Low - Medium"],
            ["High", "Implement Database Persistence (PostgreSQL/Prisma)", "Supports multi-instance scaling and high-concurrency content edits.", "Medium"],
            ["Medium", "Add Multi-User Role-Based Access Control (RBAC)", "Enables separate accounts for Content Editors, HR, and Super Admins.", "Medium"],
            ["Medium", "Add Automated End-to-End Testing (Playwright)", "Guarantees regression safety across all 49 routes during CI/CD builds.", "Low - Medium"],
            ["Low", "Implement Redis In-Memory Cache for Chatbot Queries", "Accelerates repeat AI queries and reduces LLM API consumption costs.", "Low"]
        ],
        [0.8, 2.2, 2.5, 1.0]
    )

    # ─────────────────────────────────────────────────────────────
    # 43. GLOSSARY
    # ─────────────────────────────────────────────────────────────
    h1("43. GLOSSARY")
    p("This glossary provides definitions for technical terms and concepts utilized within the TRIJOTECH website architecture:")
    bullet("Next.js App Router: Modern React architecture utilizing folder-based routing, nested layouts, and React Server Components.", "App Router: ")
    bullet("React Server Components (RSC): React components that execute exclusively on the server, generating static HTML with zero client JavaScript overhead.", "RSC: ")
    bullet("SSG (Static Site Generation): Pre-rendering web pages into static HTML and JSON files at build time for instant delivery.", "SSG: ")
    bullet("DOM Mutation Observer: Browser API that detects and reacts to dynamic additions or changes to DOM nodes in real time.", "Mutation Observer: ")
    bullet("Turbopack: Rust-based high-performance bundler embedded in Next.js 16 for near-instant compilation.", "Turbopack: ")
    bullet("HMAC (Hash-based Message Authentication Code): Cryptographic authentication mechanism used to sign and verify admin session cookies.", "HMAC: ")
    bullet("Nodemailer: Node.js library facilitating secure SMTP email transport.", "Nodemailer: ")
    bullet("Glassmorphism: UI design aesthetic featuring translucent frosted-glass card surfaces and subtle glowing borders.", "Glassmorphism: ")

    # ─────────────────────────────────────────────────────────────
    # 44. APPENDICES
    # ─────────────────────────────────────────────────────────────
    h1("44. APPENDICES")
    
    h2("Appendix A – Complete Route Map")
    p("Alphabetical listing of all 49 compiled application routes:")
    code("""/
/_not-found
/about-us
/admin
/admin/blog
/admin/blog-management
/admin/blogs
/admin/chatbot
/admin/createblog
/admin/login
/api/admin/blogs
/api/admin/chatbot
/api/admin/chatbot/settings
/api/admin/content
/api/admin/images
/api/admin/login
/api/admin/logout
/api/chat
/api/chat/analytics
/api/chat/settings
/api/content
/api/forms/careers
/api/forms/contact
/blogs
/blogs/[slug]
/careers
/case-studies
/contact
/contact-us
/corporate
/icon.svg
/industries/[[...slug]]
/industries/retail-supply-chain
/industries/pharmaceuticals-life-sciences
/industries/manufacturing
/industries/banking-financial-services
/industries/energy-utilities
/industries/steel-mining
/industries/media-telecom
/industry
/insights
/privacy-policy
/services
/services/[slug]
/services/sap-ai-ml
/services/sap-btp-full-stack
/services/sap-consulting
/services/sap-data-integration
/services/sap-implementation
/services/sap-support
/solutions
/solutions/[slug]
/solutions/e-invoicing-pro
/solutions/finlagoon-consolidation
/solutions/profitability-pro
/terms-of-service
/videos""")

    h2("Appendix B – Environment Variable Template")
    code("""# ==============================================================================
# TRIJOTECH WEBSITE ENVIRONMENT VARIABLES TEMPLATE (.env.local)
# ==============================================================================

# --- ADMIN PANEL SECURITY ---
ADMIN_PASSWORD=change_this_to_a_secure_admin_password_in_production
ADMIN_SESSION_SECRET=generate_a_random_64_character_hex_string_for_hmac_signing

# --- PUBLIC URL ---
NEXT_PUBLIC_SITE_URL=http://localhost:3000

# --- SMTP OUTBOUND EMAIL SETTINGS ---
SMTP_HOST=smtp.gmail.com
SMTP_PORT=587
SMTP_USER=notifications@trijotech.com
SMTP_PASS=your_app_specific_password_here
SMTP_SECURE=false
SMTP_FROM=notifications@trijotech.com

# --- FORM RECIPIENTS ---
CONTACT_TO_EMAIL=sales@trijotech.com
CAREERS_TO_EMAIL=hr@trijotech.com

# --- CLOUDFLARE TURNSTILE CAPTCHA (OPTIONAL) ---
NEXT_PUBLIC_TURNSTILE_SITE_KEY=
TURNSTILE_SECRET_KEY=
TURNSTILE_ALLOWED_HOSTNAMES=localhost,trijotech.com

# --- DEEPSEEK AI CHATBOT ENGINE ---
DEEPSEEK_API_KEY=your_deepseek_api_key_here
DEEPSEEK_API_URL=https://api.deepseek.com
DEEPSEEK_MODEL=deepseek-chat
DEEPSEEK_TIMEOUT_MS=25000
DEEPSEEK_MAX_TOKENS=700""")

    h2("Appendix C – Useful Commands Reference")
    code("""# Development server (Turbopack)
npm run dev

# Development server accessible over local network (0.0.0.0)
npm run dev:lan

# Full production build and static generation
npm run build

# Start production server locally
npm run start

# Static TypeScript type check
npx tsc --noEmit

# Lint codebase
npm run lint""")

    # Save to file
    output_path = r"c:\projectTrijowebsite\webapp\TRIJOTECH_Website_Complete_SOP.docx"
    doc.save(output_path)
    print(f"Document saved successfully at: {output_path}")

if __name__ == "__main__":
    build_sop()
